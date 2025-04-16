"""
Document Generator for BIM reports - Supports generating .doc and .pdf formats
"""

import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
import time
import logging
import os
from markdown import markdown

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.template_path = Path("data/templates/bim_template.md")
        self.reports_path = Path("data/reports")
        self.reports_path.mkdir(parents=True, exist_ok=True)
        
    def _load_template(self) -> str:
        """Load the BIM template."""
        with open(self.template_path, 'r') as f:
            return f.read()
            
    def _load_logs(self, service_name: str) -> List[Dict[str, Any]]:
        """Load logs for the given service."""
        log_path = Path(f"data/logs/service_logs/{service_name}_service.json")
        with open(log_path, 'r') as f:
            return json.load(f)
            
    def _analyze_logs(self, logs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze logs and extract relevant information."""
        # Get the first log entry with warnings
        log_entry = None
        warnings = []
        errors = []
        
        for entry in logs:
            log_data = json.loads(json.loads(entry['line'])['log'])
            
            # Store the first entry for service info
            if not log_entry:
                log_entry = entry
                
            # Collect warnings
            if 'webhookpost Response' in log_data and log_data['webhookpost Response'].get('warnings'):
                warnings.extend(log_data['webhookpost Response']['warnings'])
                
            # Collect errors
            if log_data.get('level') == 'error':
                errors.append(log_data)
                
        if not log_entry:
            raise ValueError("No valid log entries found")
            
        # Get service info from the first entry
        service_info = {
            'service': log_entry['fields']['app_kubernetes_io_instance'],
            'cluster': log_entry['fields']['cluster'],
            'pod': log_entry['fields']['pod'],
            'container': log_entry['fields']['container']
        }
        
        # Get metrics from the first entry with webhookpost Response
        for entry in logs:
            log_data = json.loads(json.loads(entry['line'])['log'])
            if 'webhookpost Response' in log_data:
                webhook_response = log_data['webhookpost Response']
                service_info.update({
                    'daily_bid_requests': webhook_response['daily_bid_requests'],
                    'total_bid_requests': webhook_response['total_bid_requests'],
                    'unique_devices': webhook_response['unique_devices'],
                    'unique_ips': webhook_response['unique_ips']
                })
                break
                
        service_info['warnings'] = warnings
        service_info['errors'] = errors
        return service_info
        
    def _format_warnings(self, warnings: list) -> str:
        """Format warning messages."""
        formatted = "**Warning Messages Found:**\n\n"
        for warning in warnings:
            formatted += f"- Code: {', '.join(warning['code'])}\n"
            formatted += f"  Message: {', '.join(warning['message'])}\n"
        return formatted
        
    def _format_errors(self, errors: list) -> str:
        """Format error messages."""
        if not errors:
            return "No errors found in the logs."
            
        formatted = "**Error Messages Found:**\n\n"
        for error in errors:
            formatted += f"- Level: {error['level']}\n"
            formatted += f"  Message: {error['msg']}\n"
            if 'error' in error:
                formatted += f"  Details: {error['error']}\n"
        return formatted
        
    def _convert_format(self, input_file: Path, output_format: str) -> Path:
        """Convert document to desired format without external dependencies."""
        input_file = Path(input_file)
        output_file = input_file.parent / f"{input_file.stem}.{output_format}"
        
        try:
            # If input is already in the requested format, just return it
            if input_file.suffix[1:] == output_format:
                return input_file
                
            # Read the input file content - use binary mode first to detect encoding issues
            try:
                # First try UTF-8 encoding
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with latin-1 (which accepts any byte value)
                with open(input_file, 'r', encoding='latin-1') as f:
                    content = f.read()
            
            # Handle different output formats
            if output_format == 'pdf':
                # Use WeasyPrint for PDF conversion (pure Python)
                try:
                    from weasyprint import HTML
                    HTML(string=content).write_pdf(output_file)
                    return output_file
                except ImportError:
                    logger.error("WeasyPrint not installed. Please install it with: pip install weasyprint")
                    raise
                except Exception as e:
                    logger.error(f"Error converting to PDF: {str(e)}")
                    raise
                    
            elif output_format == 'doc' or output_format == 'docx':
                # Use python-docx for DOCX creation (pure Python)
                try:
                    from docx import Document
                    from docx.shared import Inches
                    
                    # Create a new Document
                    doc = Document()
                    
                    # Add a title
                    doc.add_heading('Debug Report', 0)
                    
                    # Simple HTML parsing to extract content and ignore CSS
                    # Split into lines and process
                    in_style_block = False
                    in_pre_block = False
                    clean_content = []
                    
                    for line in content.split('\n'):
                        # Skip CSS/style content
                        if line.strip().startswith('body {') or line.strip().startswith('h1 {') or line.strip().startswith('.section {') or line.strip().startswith('pre {'):
                            in_style_block = True
                            continue
                        
                        if in_style_block and line.strip().endswith('}'):
                            in_style_block = False
                            continue
                            
                        if in_style_block:
                            continue
                            
                        # Check for pre blocks (code blocks)
                        if line.strip() == '<pre>':
                            in_pre_block = True
                            continue
                            
                        if line.strip() == '</pre>':
                            in_pre_block = False
                            continue
                        
                        # Skip most HTML tags but keep content
                        if (line.strip().startswith('<') and line.strip().endswith('>')) and not ('</' in line and '>' in line):
                            continue
                            
                        # Add line to our cleaned content
                        clean_content.append(line)
                    
                    # Now process the cleaned content
                    current_text = ""
                    
                    for line in clean_content:
                        # Try to detect headings
                        if line.strip().startswith('**') and line.strip().endswith('**'):
                            # If we have accumulated text, add it as a paragraph
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            
                            # Add the heading (remove ** markers)
                            heading_text = line.strip().replace('**', '')
                            if heading_text.strip():
                                doc.add_heading(heading_text, level=2)
                        elif line.strip().startswith('# '):
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            doc.add_heading(line[2:], level=1)
                        elif line.strip().startswith('## '):
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            doc.add_heading(line[3:], level=2)
                        elif line.strip().startswith('### '):
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            doc.add_heading(line[4:], level=3)
                        elif line.strip().startswith('- '):
                            # Handle list items
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            doc.add_paragraph(line.strip(), style='List Bullet')
                        elif line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.'):
                            # Handle numbered list items
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                            doc.add_paragraph(line.strip(), style='List Number')
                        elif not line.strip():
                            # Empty line - add paragraph break
                            if current_text.strip():
                                doc.add_paragraph(current_text)
                                current_text = ""
                        else:
                            # Add to current paragraph text
                            if current_text:
                                current_text += " " + line.strip()
                            else:
                                current_text = line.strip()
                    
                    # Add any remaining text
                    if current_text.strip():
                        doc.add_paragraph(current_text)
                    
                    # Save the document
                    doc.save(str(output_file))
                    return output_file
                except ImportError:
                    logger.error("python-docx not installed. Please install it with: pip install python-docx")
                    raise
                except Exception as e:
                    logger.error(f"Error converting to DOCX: {str(e)}")
                    raise
            else:
                logger.error(f"Unsupported output format: {output_format}")
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Document conversion failed: {str(e)}")
            # Return original file if conversion fails
            return input_file
    
    def generate_bim_doc(self, ticket_data, format='doc'):
        """
        Generate a BIM document from ticket data in either doc or pdf format.
        
        The method uses HTML as an intermediate format internally for conversion
        purposes, but only doc and pdf are supported as final output formats.
        
        Args:
            ticket_data (dict): Data from a JIRA ticket.
            format (str): Output format for the document (only 'doc' or 'pdf').
            
        Returns:
            str: Path to the generated document.
        """
        # Load template
        template_content = self._load_template()
        
        # Format content
        filled_content = template_content.format(
            ticket_id=ticket_data.get('id', 'Unknown'),
            summary=ticket_data.get('summary', 'No summary available'),
            description=ticket_data.get('description', 'No description available'),
            status=ticket_data.get('status', 'Unknown'),
            priority=ticket_data.get('priority', 'Unknown'),
            created=ticket_data.get('created', 'Unknown'),
            updated=ticket_data.get('updated', 'Unknown')
        )
        
        # Ensure output directory exists
        output_dir = self.reports_path
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        ticket_id = ticket_data.get('id', 'unknown')
        
        # Always save as HTML first for conversion
        html_file = output_dir / f"bim_report_{ticket_id}_{timestamp}.html"
        
        # Write content to html file
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(filled_content)
            
        # Validate format
        if format not in ['doc', 'pdf']:
            logger.warning(f"Unsupported format '{format}'. Defaulting to 'doc'")
            format = 'doc'
            
        # Convert to requested format if different from HTML
        if format != 'html':
            try:
                output_file = self._convert_format(html_file, format)
                return str(output_file)
            except Exception as e:
                logger.error(f"Error converting to {format}: {e}")
                return str(html_file)
        
        return str(html_file)
    
    def get_task_description(self, issue_id: str) -> str:
        """
        Get a description of the task for this agent.
        
        Args:
            issue_id: The ID of the issue to debug
            
        Returns:
            Task description string
        """
        return (f"Generate comprehensive documentation for issue {issue_id}, including analysis results, "
                f"debugging steps taken, and recommendations.") 